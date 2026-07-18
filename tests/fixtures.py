"""Test fixtures: minimal book files for testing extraction."""

import zipfile
from pathlib import Path


FIXTURES_DIR = Path(__file__).parent / "fixtures"


# ── Minimal PDF (valid structure with text) ──────────────────


def _create_minimal_pdf(path: Path):
    """Create a minimal valid PDF with visible text."""
    content = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
        b"  /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\n"
        b"endobj\n"
        b"4 0 obj\n<< /Length 180 >>\n"
        b"stream\n"
        b"BT\n/F1 24 Tf\n100 700 Td\n(\\u0413\\u043B\\u0430\\u0432\\u0430 1: \\u041D\\u0430\\u0447\\u0430\\u043B\\u043E) Tj\n0 -30 Td\n(\\u042D\\u0442\\u043E \\u043F\\u0435\\u0440\\u0432\\u044B\\u0439 \\u0430\\u0431\\u0437\\u0430\\u0446 \\u043A\\u043D\\u0438\\u0433\\u0438.) Tj\n0 -20 Td\n(\\u0412\\u0442\\u043E\\u0440\\u043E\\u0439 \\u0430\\u0431\\u0437\\u0430\\u0446 \\u0441 \\u0447\\u0438\\u0441\\u043B\\u043E\\u043C 42 \\u0438 \\u0446\\u0438\\u0444\\u0440\\u043E\\u0439 3.14.) Tj\nET\n"
        b"endstream\n"
        b"endobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n0000000058 00000 n \n0000000115 00000 n \n0000000266 00000 n \n0000000499 00000 n \n"
        b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n562\n%%EOF\n"
    )
    path.write_bytes(content)


# ── Minimal EPUB ─────────────────────────────────────────────


def _create_minimal_epub(path: Path):
    """Create a minimal valid EPUB with one chapter."""
    # Mimetype
    container = (
        '<?xml version="1.0"?>\\n'
        '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">\\n'
        '  <rootfiles>\\n'
        '    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>\\n'
        '  </rootfiles>\\n'
        '</container>\\n'
    )
    # OPF
    opf = (
        '<?xml version="1.0"?>\\n'
        '<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="book-id">\\n'
        '  <metadata>\\n'
        '    <dc:title xmlns:dc="http://purl.org/dc/elements/1.1/">Test Book</dc:title>\\n'
        '    <dc:language xmlns:dc="http://purl.org/dc/elements/1.1/">ru</dc:language>\\n'
        '  </metadata>\\n'
        '  <manifest>\\n'
        '    <item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>\\n'
        '    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>\\n'
        '  </manifest>\\n'
        '  <spine toc="ncx">\\n'
        '    <itemref idref="chapter1"/>\\n'
        '  </spine>\\n'
        '</package>\\n'
    )
    # XHTML chapter
    chapter = (
        '<?xml version="1.0" encoding="UTF-8"?>\\n'
        '<html xmlns="http://www.w3.org/1999/xhtml">\\n'
        '<head><title>Глава 1: Начало</title></head>\\n'
        '<body>\\n'
        '  <h1>Глава 1: Начало</h1>\\n'
        '  <p>Это первый абзац книги.</p>\\n'
        '  <p>Второй абзац с числом 42 и цифрой 3.14.</p>\\n'
        '</body>\\n'
        '</html>\\n'
    )
    # NCX
    ncx = (
        '<?xml version="1.0" encoding="UTF-8"?>\\n'
        '<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">\\n'
        '<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">\\n'
        '  <head>\\n'
        '    <meta name="dtb:uid" content="book-id"/>\\n'
        '  </head>\\n'
        '  <navMap>\\n'
        '    <navPoint id="chap1" playOrder="1">\\n'
        '      <navLabel><text>Глава 1: Начало</text></navLabel>\\n'
        '      <content src="chapter1.xhtml"/>\\n'
        '    </navPoint>\\n'
        '  </navMap>\\n'
        '</ncx>\\n'
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr("META-INF/container.xml", container)
        zf.writestr("OEBPS/content.opf", opf)
        zf.writestr("OEBPS/chapter1.xhtml", chapter)
        zf.writestr("OEBPS/toc.ncx", ncx)


def _create_epub_manifest_vs_spine(path: Path):
    """Create EPUB where manifest order (3,1,2) differs from spine order (1,2,3).
    
    This tests that extractor uses spine order, not manifest order.
    """
    container = (
        '<?xml version="1.0"?>'
        '<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
        '  <rootfiles>'
        '    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>'
        '  </rootfiles>'
        '</container>'
    )
    # Manifest lists chapters as 3, 1, 2 but spine says 1, 2, 3
    opf = (
        '<?xml version="1.0"?>'
        '<package xmlns="http://www.idpf.org/2007/opf" version="2.0" unique-identifier="book-id">'
        '  <metadata>'
        '    <dc:title xmlns:dc="http://purl.org/dc/elements/1.1/">Test Book</dc:title>'
        '    <dc:language xmlns:dc="http://purl.org/dc/elements/1.1/">ru</dc:language>'
        '  </metadata>'
        '  <manifest>'
        '    <item id="chapter3" href="chapter3.xhtml" media-type="application/xhtml+xml"/>'
        '    <item id="chapter1" href="chapter1.xhtml" media-type="application/xhtml+xml"/>'
        '    <item id="chapter2" href="chapter2.xhtml" media-type="application/xhtml+xml"/>'
        '    <item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>'
        '  </manifest>'
        '  <spine toc="ncx">'
        '    <itemref idref="chapter1"/>'
        '    <itemref idref="chapter2"/>'
        '    <itemref idref="chapter3"/>'
        '  </spine>'
        '</package>'
    )
    chapter1 = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<html xmlns="http://www.w3.org/1999/xhtml">'
        '<head><title>Глава 1</title></head>'
        '<body><h1>Глава 1</h1><p>Текст первой главы.</p></body>'
        '</html>'
    )
    chapter2 = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<html xmlns="http://www.w3.org/1999/xhtml">'
        '<head><title>Глава 2</title></head>'
        '<body><h1>Глава 2</h1><p>Текст второй главы.</p></body>'
        '</html>'
    )
    chapter3 = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<html xmlns="http://www.w3.org/1999/xhtml">'
        '<head><title>Глава 3</title></head>'
        '<body><h1>Глава 3</h1><p>Текст третьей главы.</p></body>'
        '</html>'
    )
    ncx = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<!DOCTYPE ncx PUBLIC "-//NISO//DTD ncx 2005-1//EN" "http://www.daisy.org/z3986/2005/ncx-2005-1.dtd">'
        '<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">'
        '  <head><meta name="dtb:uid" content="book-id"/></head>'
        '  <navMap>'
        '    <navPoint id="chap1" playOrder="1"><navLabel><text>Глава 1</text></navLabel>'
        '<content src="chapter1.xhtml"/></navPoint>'
        '    <navPoint id="chap2" playOrder="2"><navLabel><text>Глава 2</text></navLabel>'
        '<content src="chapter2.xhtml"/></navPoint>'
        '    <navPoint id="chap3" playOrder="3"><navLabel><text>Глава 3</text></navLabel>'
        '<content src="chapter3.xhtml"/></navPoint>'
        '  </navMap>'
        '</ncx>'
    )
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("mimetype", "application/epub+zip")
        zf.writestr("META-INF/container.xml", container)
        zf.writestr("OEBPS/content.opf", opf)
        zf.writestr("OEBPS/chapter1.xhtml", chapter1)
        zf.writestr("OEBPS/chapter2.xhtml", chapter2)
        zf.writestr("OEBPS/chapter3.xhtml", chapter3)
        zf.writestr("OEBPS/toc.ncx", ncx)


# ── Minimal DOCX ─────────────────────────────────────────────


def _create_minimal_docx(path: Path):
    """Create a minimal valid DOCX with headings and text."""
    # Simple approach: use python-docx if available, else create XML manually
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Глава 1: Начало", level=1)
        doc.add_paragraph("Это первый абзац книги.")
        doc.add_paragraph("Второй абзац с числом 42 и цифрой 3.14.")
        doc.save(str(path))
        return
    except ImportError:
        pass

    # Manual XML creation for minimal DOCX — too complex, rely on python-docx
    raise RuntimeError("python-docx required for test fixtures. Run: pip install python-docx")


# ── Public API ───────────────────────────────────────────────


def ensure_fixtures():
    """Create all test fixtures. Call once before tests."""
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)

    pdf_path = FIXTURES_DIR / "sample.pdf"
    if not pdf_path.exists():
        _create_minimal_pdf(pdf_path)

    epub_path = FIXTURES_DIR / "sample.epub"
    if not epub_path.exists():
        _create_minimal_epub(epub_path)

    try:
        docx_path = FIXTURES_DIR / "sample.docx"
        if not docx_path.exists():
            _create_minimal_docx(docx_path)
    except (ImportError, RuntimeError):
        pass  # DOCX tests will be skipped

    return FIXTURES_DIR


# ── Get paths ────────────────────────────────────────────────


def get_fixtures() -> dict[str, Path]:
    """Return dict of format -> Path for available fixtures."""
    fixtures = {}
    for f in FIXTURES_DIR.glob("sample.*"):
        fixtures[f.suffix] = f
    return fixtures
